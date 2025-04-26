# Импорт необходимых библиотек
import asyncio  # Для асинхронного программирования
import imaplib  # Для работы с IMAP протоколом
import email  # Для работы с email сообщениями
import requests  # Для HTTP запросов
import json  # Для работы с JSON данными
from datetime import datetime, timedelta  # Для работы с датами и временем
from docx import Document  # Для работы с Word документами
from aiogram import Bot, Dispatcher, types  # Для работы с Telegram ботом
from aiogram.fsm.storage.memory import MemoryStorage  # Для хранения состояний в памяти
import yaml  # Для работы с YAML конфигурацией
from aiogram.filters import Command  # Для обработки команд
from pprint import pprint  # Для красивого вывода данных
import smtplib  # Для работы с SMTP протоколом
from email.mime.text import MIMEText  # Для создания email сообщений
import os  # Для работы с файловой системой
import logging  # Для логирования
from aiogram.types import FSInputFile  # Добавляем импорт в начало файла

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('bot.log'),  # Логи в файл
        logging.StreamHandler()  # Логи в консоль
    ]
)

# Загрузка конфигурации из файла
with open('config.yaml', 'r', encoding='utf-8') as file:
    config = yaml.safe_load(file)

# Инициализация переменных из конфигурации
TOKEN = config['telegram']['token']  # Токен Telegram бота
BOT_EMAIL = config['email']['bot_email']  # Email бота
BOT_EMAIL_PASSWORD = config['email']['bot_password']  # Пароль от email бота
IMAP_SERVER = config['email']['imap_server']  # IMAP сервер
TRUETABS_TOKEN = config['truetabs']['token']  # Токен для API TrueTabs
TRUETABS_BASE_URL = config['truetabs']['base_url']  # Базовый URL API TrueTabs
REPORT_OUTPUT_FILE = config['report']['output_file']  # Файл для сохранения отчетов
USER_STATE_FILE = config['report']['user_state_file']  # Файл для сохранения состояний пользователей

# Инициализация бота и диспетчера
bot = Bot(token=TOKEN)  # Создание экземпляра бота
dp = Dispatcher(storage=MemoryStorage())  # Создание диспетчера с хранилищем в памяти

# Функция для загрузки состояний пользователей из файла
def load_user_states():
    """
    Загружает состояния пользователей из JSON файла.
    
    Returns:
        dict: Словарь с состояниями пользователей или пустой словарь в случае ошибки
    """
    try:
        if not os.path.exists(USER_STATE_FILE):  # Проверка существования файла
            return {}  # Возвращаем пустой словарь, если файл не существует
        
        with open(USER_STATE_FILE, 'r', encoding='utf-8') as f:
            content = f.read()
            if not content:  # Проверка на пустой файл
                return {}
            
            user_states = json.loads(content)  # Загрузка JSON данных
            
            # Преобразование строковых дат в объекты datetime
            for user_id in user_states:
                if "last_reminders" in user_states[user_id]:
                    reminders = user_states[user_id]["last_reminders"]
                    for project_name in reminders:
                        if reminders[project_name]:
                            reminders[project_name] = datetime.fromisoformat(reminders[project_name])
            
            return user_states
            
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logging.error(f"Ошибка при загрузке состояний пользователей: {e}")
        return {}

# Функция для сохранения состояний пользователей в файл
def save_user_states(user_states):
    """
    Сохраняет состояния пользователей в JSON файл.
    
    Args:
        user_states (dict): Словарь с состояниями пользователей
    """
    states_to_save = {}  # Создаем копию для сохранения
    
    # Преобразуем данные для сохранения
    for user_id, state in user_states.items():
        states_to_save[str(user_id)] = {
            "access": state.get("access", False),
            "email": state.get("email", None),
            "email_password": state.get("email_password", None),
            "smtp_server": state.get("smtp_server", None),
            "smtp_port": state.get("smtp_port", None),
            "last_reminders": {}
        }
        
        # Преобразуем datetime в строки для сохранения
        if "last_reminders" in state:
            for project_name, reminder_time in state["last_reminders"].items():
                if isinstance(reminder_time, datetime):
                    states_to_save[str(user_id)]["last_reminders"][project_name] = reminder_time.isoformat()
                else:
                    states_to_save[str(user_id)]["last_reminders"][project_name] = None
    
    # Сохраняем данные в файл
    with open(USER_STATE_FILE, 'w', encoding='utf-8') as f:
        json.dump(states_to_save, f, ensure_ascii=False, indent=2)

# Инициализация состояний пользователей
user_states = load_user_states()  # Загружаем состояния при запуске

# Обработчик команды /start
@dp.message(Command("start"))
async def start_command(message: types.Message):
    """
    Обрабатывает команду /start и инициализирует нового пользователя.
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    try:
        user_id = str(message.from_user.id)

        # Инициализация состояния нового пользователя
        if user_id not in user_states:
            user_states[user_id] = {
                "access": True,
                "email": None,
                "email_password": None,
                "smtp_server": None,
                "smtp_port": None,
                "imap_server": None,
                "imap_port": None,
                "last_reminders": {}
            }
            save_user_states(user_states)

        # Отправляем приветственное сообщение
        await message.answer(
            "🤖 Бот запущен!\n"
            "Для настройки почты используйте следующие команды:\n"
            "1️⃣ Установить email и пароль (сначала выполняете эту команду):\n"
            "/set_credentials <email> <пароль>\n"
            "2️⃣ Установить SMTP сервер и порт:\n"
            "/set_smtp <сервер> <порт>\n"
            "3️⃣ Установить IMAP сервер и порт:\n"
            "/set_imap <сервер> <порт>\n"
            "4️⃣ Проверить настройки почты (вызывайте эту команду сразу после подключения smtp и imap серверов):\n"
            "/check_settings\n"
            "📋 Другие команды:\n"
            "📊 /generate_report - Сгенерировать отчет\n"
            "📝 /send_a_form - Отправить форму для отклика\n"
            "✔️ /status - Показывает текущие настройки пользователя\n"
            "⏳ /deadlines - Просмотреть горящие дедлайны\n"
            "ℹ️ Проверка дедлайнов и новых писем выполняется автоматически"
        )

        # Если пользователь уже настроил почту, можно сразу проверить дедлайны
        await check_deadline()
    except Exception as e:
        logging.error(f"Ошибка в команде start: {e}")
        await message.answer("Произошла ошибка при запуске бота. Пожалуйста, попробуйте еще раз.")


# Новая команда /deadlines для просмотра горящих дедлайнов
@dp.message(Command("deadlines"))
async def show_deadlines(message: types.Message):
    """
    Показывает все горящие дедлайны (меньше или равно 3 дням).
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    try:
        # Получаем дедлайны
        projects = await get_deadlines()
        if not projects:
            await message.answer("✅ Нет активных проектов с близкими дедлайнами.")
            return

        current_date = datetime.now().date()
        urgent_projects = []

        for project_name, current_spending, deadline in projects:
            try:
                # Преобразуем дату дедлайна
                if isinstance(deadline, int):
                    deadline_date = datetime.fromtimestamp(deadline / 1000)
                else:
                    try:
                        deadline_date = datetime.strptime(deadline, '%Y-%m-%d')
                    except ValueError:
                        logging.error(f"Неверный формат даты для проекта {project_name}: {deadline}")
                        continue

                days_left = (deadline_date - datetime.now()).days
                if 0 <= days_left <= 3:  # Горящие дедлайны
                    responsible_persons = await get_responsible_persons_for_project(project_name)
                    if days_left == 0:
                        urgent_projects.append(
                            f"⚠️ ВНИМАНИЕ! Дедлайн проекта '{project_name}' сегодня ({deadline_date.strftime('%d.%m.%Y')})!"
                        )
                    else:
                        urgent_projects.append(
                            f"⚠️ ВНИМАНИЕ! Дедлайн проекта '{project_name}' наступает через {days_left} дней ({deadline_date.strftime('%d.%m.%Y')})!"
                        )
            except Exception as e:
                logging.error(f"Ошибка при обработке проекта {project_name}: {e}")
                continue

        if not urgent_projects:
            await message.answer("✅ Нет активных проектов с близкими дедлайнами.")
            return

        # Формируем сообщение с горящими дедлайнами
        response = "🔥 Горящие дедлайны:\n\n" + "\n\n".join(urgent_projects)
        await message.answer(response)
    except Exception as e:
        logging.error(f"Ошибка в команде deadlines: {e}")
        await message.answer("❌ Произошла ошибка при получении дедлайнов.")

async def check_deadline():
    """
    Проверяет дедлайны проектов и отправляет уведомления пользователям раз в день.
    """
    try:
        projects = await get_deadlines()  # Получаем дедлайны
        if not projects:
            logging.info("Нет проектов для проверки")
            return

        current_date = datetime.now().date()

        for project_name, current_spending, deadline in projects:
            try:
                # Преобразуем timestamp в datetime
                if isinstance(deadline, int):
                    deadline_date = datetime.fromtimestamp(deadline / 1000)
                else:
                    try:
                        deadline_date = datetime.strptime(deadline, '%Y-%m-%d')
                    except ValueError:
                        logging.error(f"Неверный формат даты для проекта {project_name}: {deadline}")
                        continue

                days_left = (deadline_date - datetime.now()).days
                if 0 <= days_left <= 3:  # Если дедлайн через 3 дня или меньше
                    # Получаем ответственных лиц
                    responsible_persons = await get_responsible_persons_for_project(project_name)

                    # Формируем сообщение с новым шаблоном
                    if days_left == 0:
                        message = (
                            f"⚠️ ВНИМАНИЕ! Дедлайн проекта '{project_name}' сегодня ({deadline_date.strftime('%d.%m.%Y')})!"
                        )
                    else:
                        message = (f"⚠️ ВНИМАНИЕ! Дедлайн проекта '{project_name}' наступает через {days_left} дней ({deadline_date.strftime('%d.%m.%Y')})!"
                        )

                    # Отправляем сообщение всем пользователям с доступом
                    for user_id, state in user_states.items():
                        if state.get('access', False):
                            # Проверяем, было ли уже отправлено уведомление сегодня
                            last_notification = state.get("last_reminders", {}).get(project_name)
                            # Отправляем уведомление если:
                            # 1. Уведомлений еще не было
                            # 2. Последнее уведомление было в другой день
                            if not last_notification or last_notification.date() < current_date:
                                try:
                                    await bot.send_message(user_id, message)
                                    user_states[user_id]["last_reminders"][project_name] = datetime.now()
                                    save_user_states(user_states)
                                except Exception as e:
                                    logging.error(f"Ошибка при отправке уведомления пользователю {user_id}: {e}")
            except Exception as e:
                logging.error(f"Ошибка при обработке проекта {project_name}: {e}")
                continue
    except Exception as e:
        logging.error(f"Ошибка в функции check_deadline: {e}")


async def get_responsible_persons_for_project(project_name: str) -> list:
    """
    Получает список ответственных лиц для проекта из API TrueTabs.
    Args:
        project_name (str): Название проекта
    Returns:
        list: Список имен ответственных лиц
    """
    try:
        # Шаг 1: Получаем данные о проектах
        projects = await fetch_data(
            "dstWYwyHfv92fusEup",  # ID таблицы "Проекты"
            "viwTPHv90rDYx",      # ID представления таблицы "Проекты"
            ["Название проекта", "Сотрудники в команде"]  # Поля для запроса
        )

        if not projects:
            logging.info("Нет проектов для проверки")
            return []

        # Находим recordIds ответственных для указанного проекта
        responsible_record_ids = []
        for project in projects:
            if project.get('fields', {}).get('Название проекта') == project_name:
                responsible_record_ids = project['fields'].get('Сотрудники в команде', [])
                break

        if not responsible_record_ids:
            logging.info(f"Для проекта '{project_name}' не найдено ответственных лиц")
            return []

        # Шаг 2: Получаем ФИО сотрудников по их recordId
        employee_map = await get_employee_details(responsible_record_ids)

        # Формируем список ответственных лиц
        responsible_persons = [employee_map.get(record_id) for record_id in responsible_record_ids]
        responsible_persons = [name for name in responsible_persons if name]  # Убираем None

        return responsible_persons

    except Exception as e:
        logging.error(f"Ошибка при получении ответственных лиц для проекта {project_name}: {e}")
        return []

async def get_employee_details(record_ids):
    """
    Получает детальные данные о сотрудниках по их recordId.
    Args:
        record_ids (list): Список recordId сотрудников
    Returns:
        dict: Отображение recordId -> ФИО сотрудника
    """
    try:
        # Параметры запроса
        headers = {'Authorization': f'Bearer {TRUETABS_TOKEN}'}
        params = {
            'viewId': 'viw5Kn9wvip4E',  # ID представления таблицы "Сотрудники"
            'filterByFormula': f'recordId() IN ({",".join(f"'{id}'" for id in record_ids)})',
            'fields': ['ФИО']
        }

        # Делаем запрос к API
        response = requests.get(
            f"{TRUETABS_BASE_URL}/fusion/v1/datasheets/dstbuj1jk5ZgFAP3V1/records",
            headers=headers,
            params=params
        )

        if response.status_code != 200:
            logging.error(f"Ошибка при запросе к API: {response.status_code} - {response.text}")
            return {}

        data = response.json()
        employees = data.get('items', [])

        # Создаем отображение recordId -> ФИО
        employee_map = {emp['recordId']: emp['fields'].get('ФИО', 'Не указано') for emp in employees}

        return employee_map

    except Exception as e:
        logging.error(f"Ошибка при получении данных о сотрудниках: {e}")
        return {}


@dp.message(Command("status"))
async def status_command(message: types.Message):
    """
    Показывает текущие настройки пользователя.
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    try:
        user_id = str(message.from_user.id)
        state = user_states.get(user_id)

        if not state:
            await message.answer("❌ Настройки не найдены. Пожалуйста, выполните команду /start.")
            return

        email = state.get("email", "Не установлен")
        smtp_server = state.get("smtp_server", "Не установлен")
        smtp_port = state.get("smtp_port", "Не установлен")
        imap_server = state.get("imap_server", "Не установлен")
        imap_port = state.get("imap_port", "Не установлен")

        await message.answer(
            f"⚙️ Текущие настройки:\n"
            f"📧 Email: {email}\n"
            f"📤 SMTP сервер: {smtp_server}:{smtp_port}\n"
            f"📥 IMAP сервер: {imap_server}:{imap_port}"
        )
    except Exception as e:
        logging.error(f"Ошибка в команде status: {e}")
        await message.answer("❌ Произошла ошибка при получении статуса.")


# Добавьте запрос IMAP данных
@dp.message(Command("set_imap"))
async def set_imap(message: types.Message):
    """
    Обрабатывает команду /set_imap для установки IMAP сервера и порта.
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    try:
        args = message.text.split(maxsplit=2)
        if len(args) < 3:
            await message.answer(
                "⚠️ Формат команды: /set_imap:\n"
                "Отправьте в одном сообщении:\n"
                "(например) /set_imap imap.example.com 993"
            )
            return

        imap_server = args[1].strip().strip('<>')  # Удаляем возможные <>
        try:
            imap_port = int(args[2].strip().strip('<>'))
        except ValueError:
            await message.answer("❌ Порт должен быть числом (например, 993)")
            return

        user_id = str(message.from_user.id)

        # Инициализация состояния если нужно
        if user_id not in user_states:
            user_states[user_id] = {
                "access": True,
                "email": None,
                "email_password": None,
                "smtp_server": None,
                "smtp_port": None,
                "imap_server": None,
                "imap_port": None,
                "last_reminders": {}
            }

        # Проверка существующих настроек
        current_imap_server = user_states[user_id].get("imap_server")
        current_imap_port = user_states[user_id].get("imap_port")

        if current_imap_server and current_imap_port:
            await message.answer(
                f"ℹ️ У вас уже установлены настройки IMAP:\n"
                f"🌐 Сервер: {current_imap_server}\n"
                f"📍 Порт: {current_imap_port}\n"
                f"Хотите перезаписать их? Отправьте 'да' или 'нет'."
            )
            # Сохраняем временные данные для подтверждения
            user_states[user_id]["temp_data"] = {
                "new_imap_server": imap_server,
                "new_imap_port": imap_port,
                "action": "set_imap"
            }
            save_user_states(user_states)
            return

        # Если настроек нет, просто сохраняем новые
        user_states[user_id]["imap_server"] = imap_server
        user_states[user_id]["imap_port"] = imap_port
        save_user_states(user_states)
        await message.answer(
            f"✅ IMAP сервер и порт успешно установлены!\n"
            f"Сервер: {imap_server}\n"
            f"Порт: {imap_port}"
        )
    except Exception as e:
        logging.error(f"Ошибка в команде set_imap: {e}")
        await message.answer(f"⚠️ Критическая ошибка: {str(e)}")

# Обработчик команды /send_a_form
@dp.message(Command("send_a_form"))
async def send_form_command(message: types.Message):
    """
    Запрашивает имя и фамилию для отправки формы отклика.
    """
    try:
        user_id = str(message.from_user.id)

        # Сохраняем состояние, что пользователь начал заполнение формы
        if user_id not in user_states:
            user_states[user_id] = {
                "access": True,
                "temp_data": {}
            }

        user_states[user_id]["temp_data"] = {"waiting_for_name": True}
        save_user_states(user_states)

        await message.answer(
            "📝 Пожалуйста, введите Имя и Фамилию кандидата в формате:\n"
            "<b>Имя Фамилия</b>\n\n"
            "Например: <i>Иван Иванов</i>",
            parse_mode="HTML"
        )

    except Exception as e:
        logging.error(f"Ошибка в команде send_a_form: {e}")
        await message.answer("Произошла ошибка при обработке команды. Пожалуйста, попробуйте еще раз.")


# Обработчик ввода имени и фамилии
@dp.message(lambda message: str(message.from_user.id) in user_states and
                            user_states[str(message.from_user.id)].get("temp_data", {}).get("waiting_for_name", False))
async def process_name_input(message: types.Message):
    try:
        user_id = str(message.from_user.id)
        full_name = message.text.strip()

        # Проверяем, что введены и имя и фамилия
        if len(full_name.split()) < 2:
            await message.answer("❌ Пожалуйста, введите и Имя и Фамилию через пробел.")
            return

        # Формируем сообщение с ссылкой
        response_text = (
            f"{full_name}, Вы проявили интерес к нашей компании. "
            f"Для прохождения на собеседование заполните [форму](https://true.tabs.sale/share/shrLoN6LlnbLKjxVLYE4Z).\n\n"
        )

        # Отправляем сообщение с Markdown-разметкой для ссылки
        await message.answer(
            response_text,
            parse_mode="Markdown",
            disable_web_page_preview=False
        )

        # Сбрасываем состояние
        user_states[user_id]["temp_data"] = {}
        save_user_states(user_states)

    except Exception as e:
        logging.error(f"Ошибка при обработке имени: {e}")
        await message.answer(
            "Произошла ошибка при обработке ваших данных. Пожалуйста, попробуйте снова командой /send_a_form.")

# Функция проверки настроек почты
async def check_email_settings(user_id):
    """
    Проверяет корректность настроек почты пользователя.
    
    Args:
        user_id (str): ID пользователя в Telegram
        
    Returns:
        tuple: (bool, str) - Успешность проверки и сообщение о результате
    """
    try:
        state = user_states.get(str(user_id))
        if not state:
            return False, "Не найдены настройки пользователя"
        
        # Получаем настройки из состояния
        email = state.get("email")
        password = state.get("email_password")
        smtp_server = state.get("smtp_server")
        smtp_port = state.get("smtp_port")
        
        # Проверяем наличие всех необходимых настроек
        if not all([email, password, smtp_server, smtp_port]):
            return False, "Не все настройки почты установлены"
        
        # Проверяем SMTP подключение
        try:
            with smtplib.SMTP_SSL(smtp_server, smtp_port) if smtp_port == 465 else smtplib.SMTP(smtp_server, smtp_port) as server:
                if smtp_port != 465:
                    server.starttls()
                server.login(email, password)
                logging.info(f"✅ SMTP подключение успешно для пользователя {user_id}")
        except Exception as e:
            return False, f"Ошибка SMTP: {e}"
        
        # Проверяем IMAP подключение
        try:
            mail = imaplib.IMAP4_SSL(IMAP_SERVER)
            mail.login(email, password)
            mail.logout()
            logging.info(f"✅ IMAP подключение успешно для пользователя {user_id}")
        except Exception as e:
            return False, f"Ошибка IMAP: {e}"
        
        return True, "Все подключения работают"
        
    except Exception as e:
        return False, f"Неожиданная ошибка: {e}"

# Обработчик команды /set_credentials
@dp.message(Command("set_credentials"))
async def set_credentials(message: types.Message):
    """
    Обрабатывает команду /set_credentials для установки email и пароля.
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    try:
        args = message.text.split(maxsplit=2)
        if len(args) < 3:
            await message.answer(
                "⚠️ Пожалуйста, укажите email и пароль после команды в одном сообщении:\n"
                "/set_credentials email пароль"
            )
            return

        email_address = args[1].strip()
        password = args[2].strip()
        user_id = str(message.from_user.id)

        # Инициализация состояния если нужно
        if user_id not in user_states:
            user_states[user_id] = {
                "access": True,
                "email": None,
                "email_password": None,
                "smtp_server": None,
                "smtp_port": None,
                "imap_server": None,
                "imap_port": None,
                "last_reminders": {}
            }

        # Проверка существующих настроек
        current_email = user_states[user_id].get("email")
        current_password = user_states[user_id].get("email_password")

        if current_email and current_password:
            await message.answer(
                f"ℹ️ У вас уже установлены учетные данные:\n"
                f"📧 Email: {current_email}\n"
                f"🔑 Пароль: {'*' * len(current_password)}\n"
                f"Хотите перезаписать их? Отправьте 'да' или 'нет'."
            )
            # Сохраняем временные данные для подтверждения
            user_states[user_id]["temp_data"] = {
                "new_email": email_address,
                "new_password": password,
                "action": "set_credentials"
            }
            save_user_states(user_states)
            return

        # Если настроек нет, просто сохраняем новые
        user_states[user_id]["email"] = email_address
        user_states[user_id]["email_password"] = password
        save_user_states(user_states)
        await message.answer(
            "✅ Email и пароль успешно сохранены!\n"
            "Теперь установите SMTP сервер и порт командой:\n"
            "/set_smtp в формате: сервер порт"
        )
        # Удалим сообщение с паролем для безопасности
        await message.delete()
    except Exception as e:
        logging.error(f"Ошибка в команде set_credentials: {e}")
        await message.answer("Произошла ошибка при установке учетных данных. Пожалуйста, попробуйте еще раз.")

# Обработчик команды /check_settings
@dp.message(Command("check_settings"))
async def check_settings(message: types.Message):
    """Обрабатывает команду проверки настроек с прогресс-сообщениями"""
    try:
        user_id = str(message.from_user.id)

        # Прогресс-сообщения
        progress_msg = await message.answer("🔄 Проверяю настройки почты...")

        # Проверка SMTP
        await progress_msg.edit_text("🔌 Проверяю SMTP подключение...")
        smtp_ok, smtp_msg = await check_smtp_connection(
            user_states[user_id]["email"],
            user_states[user_id]["email_password"],
            user_states[user_id]["smtp_server"],
            user_states[user_id]["smtp_port"]
        )

        if not smtp_ok:
            await progress_msg.edit_text(f"❌ SMTP ошибка: {smtp_msg}")
            return

        # Проверка IMAP
        await progress_msg.edit_text("📨 Проверяю IMAP подключение...")
        imap_ok, imap_msg = await check_imap_connection(
            user_states[user_id]["email"],
            user_states[user_id]["email_password"]
        )

        if not imap_ok:
            await progress_msg.edit_text(f"❌ IMAP ошибка: {imap_msg}")
            return

        # Успешный результат
        await progress_msg.edit_text(
            "✅ Все подключения работают!\n\n"
            f"SMTP: {user_states[user_id]['smtp_server']}:{user_states[user_id]['smtp_port']}\n"
            f"IMAP: {IMAP_SERVER}"
        )

    except Exception as e:
        logging.exception("Ошибка в команде check_settings")
        await message.answer(f"⚠️ Произошла ошибка: {str(e)}")

# Обработчик команды /set_smtp
@dp.message(Command("set_smtp"))
async def set_smtp(message: types.Message):
    """Обрабатывает команду установки SMTP с таймаутами и асинхронной проверкой"""
    try:
        # Парсинг аргументов
        args = message.text.split(maxsplit=2)
        if len(args) < 3:
            await message.answer(
                "⚠️ Формат команды: /set_smtp:\n"
                "Отправьте в одном сообщении: \n(например) /set_smtp smtp.example.com 465"
            )
            return
        smtp_server = args[1].strip().strip('<>')  # Удаляем возможные <>
        try:
            smtp_port = int(args[2].strip().strip('<>'))
        except ValueError:
            await message.answer("❌ Порт должен быть числом (например, 465 или 587)")
            return
        user_id = str(message.from_user.id)

        # Инициализация состояния если нужно
        if user_id not in user_states:
            user_states[user_id] = {
                "access": True,
                "email": None,
                "email_password": None,
                "smtp_server": None,
                "smtp_port": None,
                "last_reminders": {}
            }

        # Проверка существующих настроек
        current_smtp_server = user_states[user_id].get("smtp_server")
        current_smtp_port = user_states[user_id].get("smtp_port")

        if current_smtp_server and current_smtp_port:
            await message.answer(
                f"ℹ️ У вас уже установлены настройки SMTP:\n"
                f"🌐 Сервер: {current_smtp_server}\n"
                f"📍 Порт: {current_smtp_port}\n"
                f"Хотите перезаписать их? Отправьте 'да' или 'нет'."
            )
            # Сохраняем временные данные для подтверждения
            user_states[user_id]["temp_data"] = {
                "new_smtp_server": smtp_server,
                "new_smtp_port": smtp_port,
                "action": "set_smtp"
            }
            save_user_states(user_states)
            return

        # Если настроек нет, просто сохраняем новые
        user_states[user_id]["smtp_server"] = smtp_server
        user_states[user_id]["smtp_port"] = smtp_port
        save_user_states(user_states)
        await message.answer(
            f"✅ SMTP сервер и порт успешно установлены!\n"
            f"Сервер: {smtp_server}\n"
            f"Порт: {smtp_port}"
            "Теперь установите IMAP сервер и порт командой:\n"
            "/set_imap в формате: сервер порт"
        )
    except Exception as e:
        logging.exception("Ошибка в команде set_smtp")
        await message.answer(f"⚠️ Критическая ошибка: {str(e)}")


async def check_imap_connection_async(email: str, password: str, imap_server: str, imap_port: int) -> tuple[bool, str]:
    """
    Асинхронная проверка IMAP подключения.
    Args:
        email (str): Email пользователя
        password (str): Пароль пользователя
        imap_server (str): IMAP сервер
        imap_port (int): Порт IMAP сервера
    Returns:
        tuple[bool, str]: Результат проверки и сообщение об ошибке (если есть)
    """
    try:
        # Создаем соединение в отдельном потоке
        def test_connection():
            try:
                # Подключаемся к IMAP серверу
                mail = imaplib.IMAP4_SSL(imap_server, imap_port, timeout=10)
                # Пытаемся войти
                mail.login(email, password)
                # Проверяем доступные папки
                status, folders = mail.list()
                if status != "OK":
                    raise Exception("Не удалось получить список папок")
                mail.logout()
                return True, ""
            except Exception as e:
                return False, str(e)

        # Запускаем проверку в отдельном потоке
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, test_connection)
    except Exception as e:
        return False, f"Неожиданная ошибка: {str(e)}"

async def check_smtp_connection_async(email: str, password: str, smtp_server: str, smtp_port: int) -> tuple[bool, str]:
    """Асинхронная проверка SMTP подключения"""
    try:
        # Создаем соединение в отдельном потоке
        def test_connection():
            try:
                if smtp_port == 465:
                    with smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=10) as server:
                        server.login(email, password)
                else:
                    with smtplib.SMTP(smtp_server, smtp_port, timeout=10) as server:
                        server.starttls()
                        server.login(email, password)
                return True, ""
            except Exception as e:
                return False, str(e)

        # Запускаем в отдельном потоке
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, test_connection)

    except Exception as e:
        return False, f"Неожиданная ошибка: {str(e)}"

# Получение дедлайнов из TrueTabs
async def get_deadlines():
    """
    Получает дедлайны проектов из API TrueTabs.
    
    Returns:
        list: Список дедлайнов проектов или пустой список в случае ошибки
    """
    headers = {
        'Authorization': f'Bearer {TRUETABS_TOKEN}',
        'Content-Type': 'application/json'
    }
    
    params = {
        'viewId': 'viwTPHv90rDYx',
        'pageNum': 1,
        'pageSize': 100
    }
    
    try:
        response = requests.get(
            f"{TRUETABS_BASE_URL}/fusion/v1/datasheets/dstWYwyHfv92fusEup/records",
            headers=headers,
            params=params
        )
        response.raise_for_status()  # Проверка на успешный ответ
        data = response.json()
        
        if 'data' not in data:
            logging.error("Ошибка: в ответе отсутствует поле 'data'")
            logging.error(f"Полный ответ: {data}")
            return []
            
        if 'records' not in data['data']:
            logging.error("Ошибка: в ответе отсутствует поле 'records'")
            return []
            
        deadlines = []
        for record in data['data']['records']:
            fields = record.get('fields', {})
            project_name = fields.get('Название проекта')
            current_spending = fields.get('Текущие затраты по проекту')
            deadline = fields.get('Дата окончания')
            
            if project_name and deadline:
                deadlines.append((project_name, current_spending, deadline))
        
        return deadlines
        
    except requests.exceptions.RequestException as e:
        logging.error(f"Ошибка при запросе к API: {e}")
        return []
    except json.JSONDecodeError as e:
        logging.error(f"Ошибка при разборе JSON: {e}")
        return []
    except Exception as e:
        logging.error(f"Неожиданная ошибка: {e}")
        return []

# Проверка новых писем
async def check_new_email():
    """
    Проверяет наличие новых писем в почтовом ящике и отправляет уведомления пользователям.
    """
    try:
        # Подключаемся к IMAP серверу
        mail = imaplib.IMAP4_SSL(IMAP_SERVER)
        mail.login(BOT_EMAIL, BOT_EMAIL_PASSWORD)
        mail.select('inbox')
        
        # Ищем непрочитанные письма
        _, message_numbers = mail.search(None, 'UNSEEN')
        
        if not message_numbers[0]:
            logging.info("Нет новых писем")
            return
            
        # Обрабатываем каждое письмо
        for num in message_numbers[0].split():
            try:
                _, msg_data = mail.fetch(num, '(RFC822)')
                email_body = msg_data[0][1]
                email_message = email.message_from_bytes(email_body)
                
                # Получаем информацию о письме
                subject = email_message["subject"]
                sender = email_message["from"]
                date = email_message["date"]
                
                # Формируем сообщение для Telegram
                message = (
                    f"📧 Новое письмо!\n"
                    f"От: {sender}\n"
                    f"Тема: {subject}\n"
                    f"Дата: {date}"
                )
                
                # Отправляем уведомление всем пользователям с доступом
                for user_id, state in user_states.items():
                    if state.get('access', False):
                        try:
                            await bot.send_message(chat_id=user_id, text=message)
                            logging.info(f"Отправлено уведомление о письме пользователю {user_id}")
                        except Exception as e:
                            logging.error(f"Ошибка при отправке уведомления пользователю {user_id}: {e}")
                            
            except Exception as e:
                logging.error(f"Ошибка при обработке письма: {e}")
                continue
                
        # Закрываем соединение
        mail.close()
        mail.logout()
        
    except imaplib.IMAP4.error as e:
        logging.error(f"Ошибка IMAP: {e}")
    except Exception as e:
        logging.error(f"Ошибка при проверке почты: {e}")
        try:
            mail.logout()
        except:
            pass

# Генерация отчета
@dp.message(Command("generate_report"))
async def generate_comprehensive_report(message: types.Message):
    """Генерация полного отчёта с данными из всех таблиц"""
    try:
        # Уведомляем пользователя о начале генерации
        await message.answer("⌛ Начинаю генерацию отчёта...")

        # 1. Получаем данные из всех таблиц
        employees = await fetch_data("dstbuj1jk5ZgFAP3V1", "viw5Kn9wvip4E",
                                     ["ФИО", "Должность", "Зарплата до вычета НДС", "Проекты", "Категория"])

        projects = await fetch_data("dstWYwyHfv92fusEup", "viwTPHv90rDYx",
                                    ["Название проекта", "Текущие затраты", "Сотрудники в команде"])

        vacancies = await fetch_data("dst8rGb4aS0aF9Rkx2", "viwwKq1FQM30f",
                                     ["Вакансия"])

        if not all([employees, projects, vacancies]):
            await message.answer("❌ Не удалось получить данные из одной или нескольких таблиц")
            return

        # 2. Обрабатываем данные
        vacancies_dict = {v['recordId']: v['fields']['Вакансия'] for v in vacancies if 'fields' in v}
        projects_dict = {p['recordId']: {
            'name': p['fields']['Название проекта'],
            'cost': p['fields'].get('Текущие затраты', 0)
        } for p in projects if 'fields' in p}

        report_data = []
        dismissed_data = []  # Для уволенных сотрудников
        total_salary = 0
        total_dismissed_salary = 0
        project_costs = {}

        for emp in employees:
            if 'fields' not in emp:
                continue

            emp_fields = emp['fields']

            # Проверяем статус сотрудника
            category = emp_fields.get('Категория', '')
            is_dismissed = category == 'Уволен'

            position_ids = emp_fields.get('Должность', [])
            position_id = position_ids[0] if isinstance(position_ids, list) and position_ids else position_ids

            employee = {
                'name': emp_fields.get('ФИО', 'Не указано'),
                'position': vacancies_dict.get(position_id, "Должность не указана"),
                'salary': float(emp_fields.get('Зарплата до вычета НДС', 0)),
                'projects': [],
                'dismissed': is_dismissed
            }

            # Добавляем в соответствующий список
            if is_dismissed:
                total_dismissed_salary += employee['salary']
                dismissed_data.append(employee)
            else:
                total_salary += employee['salary']
                report_data.append(employee)

            # Учитываем проекты для всех сотрудников (включая уволенных)
            for project_id in emp_fields.get('Проекты', []):
                if project_id in projects_dict:
                    project_name = projects_dict[project_id]['name']
                    employee['projects'].append(project_name)

                    if not is_dismissed:  # Учитываем затраты только для работающих
                        if project_name not in project_costs:
                            project_costs[project_name] = 0
                        project_costs[project_name] += employee['salary']

        # 3. Генерируем отчёт
        doc = Document()

        # Заголовок
        doc.add_heading('Полный аналитический отчёт', 0)
        doc.add_paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

        # Основные показатели (только для работающих)
        doc.add_heading('Ключевые метрики (работающие сотрудники)', level=1)
        doc.add_paragraph(f"Общий фонд зарплат: {total_salary:,.2f} ₽", style='IntenseQuote')
        doc.add_paragraph(f"Средняя зарплата: {total_salary / len(report_data):,.2f} ₽" if report_data else "0 ₽")

        # Сводная таблица по работающим сотрудникам
        doc.add_heading('Сотрудники (работающие)', level=1)
        if report_data:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Заголовки
            hdr = table.rows[0].cells
            hdr[0].text = 'ФИО'
            hdr[1].text = 'Должность'
            hdr[2].text = 'Зарплата (₽)'
            hdr[3].text = 'Проектов'

            # Данные
            for emp in sorted(report_data, key=lambda x: x['salary'], reverse=True):
                row = table.add_row().cells
                row[0].text = emp['name']
                row[1].text = emp['position']
                row[2].text = f"{emp['salary']:,.2f}"
                row[3].text = str(len(emp['projects']))
        else:
            doc.add_paragraph("Нет данных о работающих сотрудниках", style='IntenseQuote')

        # Детализация по работающим сотрудникам
        doc.add_heading('Участие в проектах (работающие)', level=1)
        for emp in report_data:
            doc.add_heading(emp['name'], level=2)
            doc.add_paragraph(f"Должность: {emp['position']}")
            doc.add_paragraph(f"Зарплата: {emp['salary']:,.2f} ₽")

            if emp['projects']:
                doc.add_paragraph("Проекты:", style='Heading3')
                for project in emp['projects']:
                    doc.add_paragraph(f"• {project}", style='ListBullet')
            else:
                doc.add_paragraph("Не участвует в проектах", style='IntenseQuote')

            doc.add_paragraph()

        # Затраты по проектам (только работающие)
        doc.add_heading('Затраты на проекты (работающие)', level=1)
        if project_costs:
            cost_table = doc.add_table(rows=1, cols=3)
            cost_table.style = 'Light Shading'

            # Заголовки
            hdr = cost_table.rows[0].cells
            hdr[0].text = 'Проект'
            hdr[1].text = 'Затраты (₽)'
            hdr[2].text = 'Участников'

            # Данные
            for project, cost in sorted(project_costs.items(), key=lambda x: x[1], reverse=True):
                row = cost_table.add_row().cells
                row[0].text = project
                row[1].text = f"{cost:,.2f}"
                row[2].text = str(sum(1 for emp in report_data if project in emp['projects']))
        else:
            doc.add_paragraph("Нет данных о затратах на проекты", style='IntenseQuote')

        # Раздел для уволенных сотрудников
        if dismissed_data:
            doc.add_heading('Уволенные сотрудники', level=1)
            doc.add_paragraph(f"Общий фонд зарплат уволенных: {total_dismissed_salary:,.2f} ₽", style='IntenseQuote')

            dismissed_table = doc.add_table(rows=1, cols=4)
            dismissed_table.style = 'Table Grid'

            # Заголовки
            hdr = dismissed_table.rows[0].cells
            hdr[0].text = 'ФИО'
            hdr[1].text = 'Должность'
            hdr[2].text = 'Зарплата (₽)'
            hdr[3].text = 'Проектов'

            # Данные
            for emp in sorted(dismissed_data, key=lambda x: x['salary'], reverse=True):
                row = dismissed_table.add_row().cells
                row[0].text = emp['name']
                row[1].text = emp['position']
                row[2].text = f"{emp['salary']:,.2f}"
                row[3].text = str(len(emp['projects']))

        # Сохраняем отчёт
        report_path = f"full_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(report_path)

        # 4. Отправляем отчёт
        await message.answer_document(
            FSInputFile(report_path),
            caption="📊 Полный аналитический отчёт"
        )

    except Exception as e:
        logging.exception("Ошибка генерации отчёта")
        await message.answer(f"❌ Произошла ошибка при генерации отчёта: {str(e)}")


async def fetch_data(datasheet_id, view_id, fields):
    """Универсальный метод получения данных из API"""
    try:
        headers = {'Authorization': f'Bearer {TRUETABS_TOKEN}'}
        params = {
            'viewId': view_id,
            'pageSize': 1000,
            'fields': fields,
            'fieldKey': 'name'
        }

        response = requests.get(
            f"{TRUETABS_BASE_URL}/fusion/v1/datasheets/{datasheet_id}/records",
            headers=headers,
            params=params
        )
        response.raise_for_status()

        return response.json().get('data', {}).get('records', [])

    except Exception as e:
        logging.error(f"Ошибка получения данных из {datasheet_id}: {e}")
        return None


async def send_email(to_email: str, subject: str, body: str, user_id: str) -> bool:
    """Отправляет email с подробным логированием и обработкой ошибок"""
    try:
        state = user_states.get(str(user_id))
        if not state:
            logging.error(f"Не найдены настройки для user_id {user_id}")
            return False

        email = state.get("email")
        password = state.get("email_password")
        smtp_server = state.get("smtp_server")
        smtp_port = state.get("smtp_port")

        # Валидация параметров
        if not all([email, password, smtp_server, smtp_port]):
            logging.error("Не все параметры SMTP заполнены")
            return False

        # Формируем сообщение
        msg = MIMEText(body, 'plain', 'utf-8')
        msg['Subject'] = subject
        msg['From'] = email
        msg['To'] = to_email

        # Логирование перед отправкой
        logging.info(f"Попытка отправки письма от {email} к {to_email} через {smtp_server}:{smtp_port}")

        # Подключение и отправка
        try:
            if smtp_port == 465:
                with smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=10) as server:
                    server.login(email, password)
                    server.send_message(msg)
            else:
                with smtplib.SMTP(smtp_server, smtp_port, timeout=10) as server:
                    server.starttls()
                    server.login(email, password)
                    server.send_message(msg)

            logging.info(f"Письмо успешно отправлено на {to_email}")
            return True

        except smtplib.SMTPRecipientsRefused as e:
            logging.error(f"Ошибка адреса получателя: {e}")
        except smtplib.SMTPException as e:
            logging.error(f"SMTP ошибка: {e}")
        except Exception as e:
            logging.error(f"Неожиданная ошибка при отправке: {e}", exc_info=True)

        return False

    except Exception as e:
        logging.exception(f"Критическая ошибка в send_email: {e}")
        return False

# Проверка и обновление статуса кандидатов
async def check_candidates_status():
    """
    Проверяет и обновляет статус кандидатов.
    """
    try:
        headers = {
            'Authorization': f'Bearer {TRUETABS_TOKEN}',
            'Content-Type': 'application/json'
        }
        
        response = requests.get(
            f"{TRUETABS_BASE_URL}/fusion/v1/datasheets/dstWYwyHfv92fusEup/records",
            headers=headers,
            params={
                'viewId': 'viwTPHv90rDYx',
                'pageNum': 1,
                'pageSize': 100
            }
        )
        
        if response.status_code != 200:
            logging.error(f"Ошибка при получении данных о кандидатах: {response.status_code}")
            return

        data = response.json()
        if 'data' not in data:
            logging.error("Ошибка: в ответе отсутствует поле 'data'")
            logging.error(f"Полный ответ: {data}")
            return
            
        if 'records' not in data['data']:
            logging.error("Ошибка: в ответе отсутствует поле 'records'")
            return

        candidates = data['data']['records']
        
        for candidate in candidates:
            try:
                fields = candidate.get('fields', {})
                status = fields.get('Статус кандидата')
                
                if status == 'Ожидает собеседование':
                    email = fields.get('Электронная почта')
                    name = fields.get('ФИО', 'Кандидат')
                    interview_date = fields.get('Дата собеседования')
                    interview_time = fields.get('Время собеседования')
                    
                    if all([email, interview_date, interview_time]):
                        # Отправляем email кандидату
                        subject = "Приглашение на собеседование"
                        body = f"""
                        Уважаемый(ая) {name}!

                        Приглашаем вас на собеседование, которое состоится {interview_date} в {interview_time}.

                        С уважением,
                        HR команда
                        """
                        
                        # Отправляем письмо от имени каждого пользователя с доступом
                        for user_id in user_states:
                            if user_states[user_id].get('access', False):
                                if await send_email(email, subject, body, user_id):
                                    break  # Если письмо успешно отправлено, прекращаем попытки
                                    
                    else:
                        logging.warning(f"⚠️ Неполные данные для кандидата {name}")
                        
            except Exception as e:
                logging.error(f"Ошибка при обработке кандидата: {e}")
                continue

    except requests.exceptions.RequestException as e:
        logging.error(f"Ошибка при запросе к API: {e}")
    except json.JSONDecodeError as e:
        logging.error(f"Ошибка при разборе JSON: {e}")
    except Exception as e:
        logging.error(f"Неожиданная ошибка при проверке статуса кандидатов: {e}")

# Периодическая проверка
async def periodic_check():
    """
    Периодическая проверка дедлайнов и писем с автоматическим восстановлением.
    """
    try:
        while True:
            try:
                logging.info("Начало периодической проверки...")
                
                # Проверяем дедлайны с повторными попытками
                try:
                    projects = await retry_on_error(get_deadlines, max_retries=3, delay=5)
                    if projects:
                        await retry_on_error(check_deadline, max_retries=3, delay=5)
                    else:
                        logging.info("Нет проектов для проверки дедлайнов")
                except Exception as e:
                    logging.error(f"Ошибка при проверке дедлайнов: {e}")
                
                # Проверяем почту с повторными попытками
                try:
                    await retry_on_error(check_new_email, max_retries=3, delay=5)
                except Exception as e:
                    logging.error(f"Ошибка при проверке почты: {e}")
                
                # Проверяем статус кандидатов с повторными попытками
                try:
                    await retry_on_error(check_candidates_status, max_retries=3, delay=5)
                except Exception as e:
                    logging.error(f"Ошибка при проверке статуса кандидатов: {e}")
                
                logging.info("Периодическая проверка завершена")
                
            except Exception as e:
                logging.error(f"Ошибка в периодической проверке: {e}")
            finally:
                # Ждем 10 минут перед следующей проверкой
                await asyncio.sleep(600)
    except asyncio.CancelledError:
        logging.info("Периодическая проверка остановлена пользователем")
        raise

#Проверка подключения к smtp
async def check_smtp_connection(email: str, password: str, smtp_server: str, smtp_port: int) -> tuple[bool, str]:
    """Проверяет SMTP подключение с подробным логированием"""
    try:
        logging.info(f"Попытка SMTP подключения к {smtp_server}:{smtp_port}...")

        # Определяем тип подключения
        if smtp_port == 465:
            server = smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=10)
        else:
            server = smtplib.SMTP(smtp_server, smtp_port, timeout=10)
            logging.info("Инициируем STARTTLS...")
            server.starttls()

        logging.info("Пытаемся войти...")
        server.login(email, password)
        server.quit()

        logging.info("SMTP подключение успешно")
        return True, "SMTP подключение успешно"

    except smtplib.SMTPAuthenticationError:
        error_msg = "Ошибка авторизации SMTP: неверный email или пароль"
        logging.error(error_msg)
        return False, error_msg
    except smtplib.SMTPConnectError:
        error_msg = f"Не удалось подключиться к SMTP серверу {smtp_server}:{smtp_port}"
        logging.error(error_msg)
        return False, error_msg
    except Exception as e:
        error_msg = f"SMTP ошибка: {str(e)}"
        logging.error(error_msg, exc_info=True)
        return False, error_msg


async def check_imap_connection(email: str, password: str) -> tuple[bool, str]:
    """Проверяет IMAP подключение с подробным логированием"""
    try:
        logging.info(f"Попытка IMAP подключения к {IMAP_SERVER}...")
        mail = imaplib.IMAP4_SSL(IMAP_SERVER, timeout=10)
        mail.login(email, password)
        mail.logout()
        logging.info("IMAP подключение успешно")
        return True, "IMAP подключение успешно"

    except imaplib.IMAP4.error as e:
        error_msg = f"IMAP ошибка: {str(e)}"
        logging.error(error_msg)
        return False, error_msg
    except Exception as e:
        error_msg = f"Неожиданная IMAP ошибка: {str(e)}"
        logging.error(error_msg, exc_info=True)
        return False, error_msg


@dp.message(lambda message: str(message.from_user.id) in user_states and
                            "temp_data" in user_states[str(message.from_user.id)] and
                            "action" in user_states[str(message.from_user.id)]["temp_data"])
async def handle_confirmation(message: types.Message):
    """
    Обрабатывает подтверждение пользователя на перезапись настроек.
    """
    try:
        user_id = str(message.from_user.id)
        temp_data = user_states[user_id]["temp_data"]
        action = temp_data.get("action")
        confirmation = message.text.strip().lower()

        if confirmation not in ["да", "нет"]:
            await message.answer("⚠️ Пожалуйста, отправьте 'да' или 'нет'.")
            return

        if confirmation == "нет":
            # Отмена действия
            user_states[user_id]["temp_data"] = {}
            save_user_states(user_states)
            await message.answer("❌ Действие отменено. Настройки остались без изменений.")
            return

        # Выполняем действие в зависимости от типа
        if action == "set_credentials":
            user_states[user_id]["email"] = temp_data["new_email"]
            user_states[user_id]["email_password"] = temp_data["new_password"]
            await message.answer(
                "✅ Email и пароль успешно обновлены!\n"
                "Теперь установите SMTP сервер и порт командой:\n"
                "/set_smtp в формате: сервер порт"
            )
        elif action == "set_smtp":
            user_states[user_id]["smtp_server"] = temp_data["new_smtp_server"]
            user_states[user_id]["smtp_port"] = temp_data["new_smtp_port"]
            await message.answer(
                f"✅ SMTP сервер и порт успешно обновлены!\n"
                f"Сервер: {temp_data['new_smtp_server']}\n"
                f"Порт: {temp_data['new_smtp_port']}"
            )
        elif action == "set_imap":
            user_states[user_id]["imap_server"] = temp_data["new_imap_server"]
            user_states[user_id]["imap_port"] = temp_data["new_imap_port"]
            await message.answer(
                f"✅ IMAP сервер и порт успешно обновлены!\n"
                f"Сервер: {temp_data['new_imap_server']}\n"
                f"Порт: {temp_data['new_imap_port']}"
            )

        # Очищаем временные данные
        user_states[user_id]["temp_data"] = {}
        save_user_states(user_states)
    except Exception as e:
        logging.error(f"Ошибка при обработке подтверждения: {e}")
        await message.answer("⚠️ Произошла ошибка при обработке вашего ответа.")

# Проверка подключения к IMAP
async def check_imap_connection(email: str, password: str) -> tuple[bool, str]:
    """Проверяет IMAP подключение с подробным логированием"""
    try:
        logging.info(f"Попытка IMAP подключения для {email}...")

        # Добавляем таймаут и явное указание SSL
        mail = imaplib.IMAP4_SSL(IMAP_SERVER, timeout=10)

        # Логируем перед аутентификацией
        logging.info(f"Пытаюсь войти... (сервер: {IMAP_SERVER})")

        # Явная обработка ошибки аутентификации
        try:
            mail.login(email, password)
        except imaplib.IMAP4.error as e:
            error_msg = f"Ошибка аутентификации IMAP: {str(e)}"
            if "AUTHENTICATIONFAILED" in str(e):
                error_msg += "\nВозможные причины:\n1. Неверный пароль\n2. IMAP отключён в настройках почты\n3. Требуется пароль приложения"
            logging.error(error_msg)
            return False, error_msg

        # Проверяем доступные папки
        status, folders = mail.list()
        if status == "OK":
            logging.info(f"Доступные папки: {', '.join(f.decode() for f in folders[:3])}...")
        else:
            logging.warning("Не удалось получить список папок")

        mail.logout()
        logging.info("IMAP подключение успешно")
        return True, "IMAP подключение успешно"

    except Exception as e:
        error_msg = f"Неожиданная IMAP ошибка: {str(e)}"
        logging.error(error_msg, exc_info=True)
        return False, error_msg

async def retry_on_error(func, *args, max_retries=3, delay=5):
    """
    Выполняет функцию с повторными попытками при ошибке.
    
    Args:
        func (callable): Функция для выполнения
        *args: Аргументы для функции
        max_retries (int): Максимальное количество попыток
        delay (int): Задержка между попытками
    """
    for attempt in range(max_retries):
        try:
            return await func(*args)
        except Exception as e:
            if attempt == max_retries - 1:
                logging.error(f"Ошибка после {max_retries} попыток: {e}")
                return None
            logging.error(f"Ошибка: {e}. Повторная попытка через {delay} секунд...")
            await asyncio.sleep(delay)

async def run_with_recovery():
    """
    Запускает бота с автоматическим восстановлением после ошибок.
    """
    while True:
        try:
            logging.info("Запуск бота...")
            
            # Запускаем периодическую проверку
            check_task = asyncio.create_task(periodic_check())
            
            # Запускаем бота
            logging.info("Бот запущен и готов к работе")
            await dp.start_polling(bot, skip_updates=True)
            
        except KeyboardInterrupt:
            logging.info("Бот остановлен пользователем")
            break  # Выход из цикла при нажатии Ctrl+C
        except Exception as e:
            logging.error(f"Критическая ошибка: {e}")
            logging.error("Перезапуск бота через 5 секунд...")
            try:
                await bot.close()
            except:
                pass
            await asyncio.sleep(5)
            continue

if __name__ == '__main__':
    try:
        asyncio.run(run_with_recovery())  # Запуск бота
    except KeyboardInterrupt:
        logging.info("Бот остановлен пользователем")
    finally:
        try:
            asyncio.run(bot.close())  # Закрытие бота
        except:
            pass
